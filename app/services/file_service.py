import os
import io
import json
from llama_index.core import Document
import pypdf
import docx


def process_file_in_memory(filename: str, file_bytes: bytes):
    text_content = ""
    file_stream = io.BytesIO(file_bytes)

    try:
        # 1. Xử lý PDF
        if filename.lower().endswith(".pdf"):
            pdf_reader = pypdf.PdfReader(file_stream)
            for page in pdf_reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text_content += extracted + "\n"

        # 2. Xử lý DOCX
        elif filename.lower().endswith(".docx") or filename.lower().endswith(".doc"):
            doc = docx.Document(file_stream)
            for para in doc.paragraphs:
                text_content += para.text + "\n"

        # 3. Xử lý TXT & MD
        elif (
            filename.lower().endswith(".txt")
            or filename.lower().endswith(".text")
            or filename.lower().endswith(".md")
        ):
            text_content = file_stream.read().decode("utf-8")

        # 4. Xử lý JSON
        elif filename.lower().endswith(".json"):
            json_str = file_stream.read().decode("utf-8")
            data = json.loads(json_str)
            text_content = json.dumps(data, ensure_ascii=False, indent=2)

        else:
            print(f"Định dạng file chưa hỗ trợ: {filename}")
            return []

    except Exception as e:
        print(f"Lỗi đọc file {filename}: {e}")
        return []

    if text_content.strip():
        return [Document(text=text_content, metadata={"filename": filename})]

    return []
